import io

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif", "image/tiff"}


class UnsupportedDocumentTypeError(Exception):
    pass


def extract_pages(file_bytes: bytes, mime_type: str) -> list[str]:
    """Return a list of page-level text strings. Index 0 is page 1.

    Formats without a native page concept (docx, xlsx) are returned as a single "page".
    Images are not text-extracted here — Document Analyzer sends them directly to a
    vision-capable AI provider instead of relying on OCR.
    """
    if mime_type == "application/pdf":
        return _extract_pdf(file_bytes)
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return [_extract_docx(file_bytes)]
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return [_extract_xlsx(file_bytes)]
    if mime_type == "text/plain":
        return [file_bytes.decode("utf-8", errors="ignore")]
    if mime_type in IMAGE_MIME_TYPES:
        return []

    raise UnsupportedDocumentTypeError(f"Unsupported document type: {mime_type}")


def _extract_pdf(file_bytes: bytes) -> list[str]:
    reader = PdfReader(io.BytesIO(file_bytes))
    return [page.extract_text() or "" for page in reader.pages]


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(file_bytes: bytes) -> str:
    workbook = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                parts.append(" | ".join("" if cell is None else str(cell) for cell in row))
    return "\n".join(parts)


def is_image(mime_type: str) -> bool:
    return mime_type in IMAGE_MIME_TYPES
